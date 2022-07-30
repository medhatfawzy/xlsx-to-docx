from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL


# Setting up some properties of the page
document = Document()
section = document.sections[0]
section.page_height = Cm(21) #A4 width
section.page_width = Cm(29.7) #A4 height
section.orientation = WD_ORIENT.LANDSCAPE
section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.2)


for i in range(1, 20):
    outer_cell = document.add_table(1, 1).cell(0,0)
    outer_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    heading = outer_cell.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("R" + str(i))
    heading_run.font.bold = True
    heading_run.font.size = Pt(350)
    document.add_page_break()



document.save("R.docx")

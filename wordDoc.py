from docx import Document
from docx.shared import Inches, Emu, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement

import math

from load_excel_file import load_excel_file

data = load_excel_file("record.xlsx")

def prevent_document_break(document):
    """
    https://github.com/python-openxml/python-docx/issues/245#event-621236139
    Globally prevent table cells from splitting across pages.
    """
    tags = document.element.xpath('//w:tr')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)  # Append in the new tag

# Setting up some properties of the page
document = Document()
section = document.sections[0]
section.page_width = Inches(8.27) #A4 width
section.page_height = Inches(11.69) #A4 height
section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.2)

# The properties of an item
KEYS = [
    "اسم الصنف",
    "رقم الصفحة",
    "اسم الدفتر",
    "الكمية",
    "الوحدة",
    "إحداثي التخزين"
]

# An outer table to create the tables inside of its cells
outer_table = document.add_table(math.ceil((len(data)) / 2), 2) 
outer_table.table_direction = WD_TABLE_DIRECTION.RTL
outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
outer_table.autofit = True

# looping through each row in the data and creating a table for each instance then adding inside the cell
for x, row in enumerate(data):
    # Skip if the quantity is zero
    if row[3] == 0:
        continue
    

    outer_cell = outer_table.cell(int(x/2), x % 2)
    outer_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    # The title of each card
    heading_paragraph = outer_cell.paragraphs[0]
    heading_paragraph.paragraph_format.keep_together = True
    heading_paragraph.paragraph_format.keep_with_next = True
    heading_paragraph.paragraph_format.space_before = Inches(0.2)
    heading_paragraph.paragraph_format.space_after = Inches(0)
    heading_paragraph.paragraph_format.alignment =  WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading_paragraph.add_run("كارت الصنف")

    # Setting some properties for the title
    heading_run.font.rtl = True
    heading_run.font.cs_bold = True

    # Creating the table for the properties inside a cell of the outer table
    table = outer_cell.add_table(len(KEYS), 2)
    table.style = "Light Shading"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.table_direction = WD_TABLE_DIRECTION.RTL
    table.style.font.rtl = True
    table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for i, key in enumerate(KEYS):
        table.cell(i, 0).text = key
        table.cell(i, 0).width = Inches(1)
        table.cell(i, 1).text = str(row[i])
        table.cell(i, 1).width = Inches(3)


prevent_document_break(document)
document.save("items_tags.docx")
